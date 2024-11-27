from typing import Dict, Any, Optional, List
import asyncio
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import ebooklib
from ebooklib import epub
import markdown
from bs4 import BeautifulSoup
from app.core.logging import app_logger
from app.services.ai_client import ai_client
from app.db.models import File, AnalysisRecord
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import aiohttp
import io
from docx import Document

class DocumentService:
    """文档处理服务"""
    
    async def extract_text(self, file_path: str) -> str:
        """从文件中提取文本"""
        try:
            # 处理远程URL
            if file_path.startswith(('http://', 'https://')):
                async with aiohttp.ClientSession() as session:
                    async with session.get(file_path) as response:
                        if response.status != 200:
                            return f"无法访问文档URL: {file_path}"
                        
                        # 读取文件内容
                        file_content = await response.read()
                        
                        # 根据文件类型处理
                        if file_path.endswith('.docx'):
                            # 将二进制内容转换为文档对象
                            doc = Document(io.BytesIO(file_content))
                            # 提取所有段落的文本
                            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                            return text
                        elif file_path.endswith('.pdf'):
                            return await self._extract_pdf_text_from_bytes(file_content)
                        elif file_path.endswith('.txt'):
                            return file_content.decode('utf-8')
                        else:
                            return f"不支持的远程文件类型: {file_path}"

            # 本地文件处理逻辑
            path = Path(file_path)
            suffix = path.suffix.lower()
            
            if suffix in ['.txt']:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    return await f.read()
            elif suffix in ['.pdf']:
                return await self._extract_pdf_text(file_path)
            elif suffix in ['.doc', '.docx']:
                return await self._extract_doc_text(file_path)
            else:
                return f"不支持的文件类型: {suffix}"

        except Exception as e:
            app_logger.error(f"文本提取失败: {str(e)}")
            return f"文本提取失败: {str(e)}"

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """提取PDF文本"""
        try:
            # 使用线程池执行同步操作
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._pdf_extract_worker, file_path)
        except Exception as e:
            raise ValueError(f"PDF文本提取失败: {str(e)}")

    def _pdf_extract_worker(self, file_path: Path) -> str:
        """PDF文本提取工作函数"""
        text = []
        with open(file_path, 'rb') as file:
            pdf = PdfReader(file)
            for page in pdf.pages:
                text.append(page.extract_text())
        return '\n'.join(text)

    async def _extract_docx_text(self, file_path: Path) -> str:
        """提取DOCX文本"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._docx_extract_worker, file_path)
        except Exception as e:
            raise ValueError(f"DOCX文本提取失败: {str(e)}")

    def _docx_extract_worker(self, file_path: Path) -> str:
        """DOCX文本提取工作函数"""
        doc = DocxDocument(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    async def _extract_epub_text(self, file_path: Path) -> str:
        """提取EPUB文本"""
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._epub_extract_worker, file_path)
        except Exception as e:
            raise ValueError(f"EPUB文本提取失败: {str(e)}")

    def _epub_extract_worker(self, file_path: Path) -> str:
        """EPUB文本提取工作函数"""
        book = epub.read_epub(str(file_path))
        text = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text.append(soup.get_text())
        return '\n'.join(text)

    async def _extract_markdown_text(self, file_path: Path) -> str:
        """提取Markdown文本"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()

    async def analyze_document(
        self,
        db: AsyncSession,
        file_id: str,
        query: Optional[str] = None,
        system_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """分析文档内容"""
        try:
            # 获取文件记录
            file_query = select(File).where(File.file_id == file_id)
            result = await db.execute(file_query)
            file_record = result.scalar_one_or_none()
            
            if not file_record:
                raise ValueError(f"File not found: {file_id}")
                
            # 获取文件路径
            file_path = Path(settings.UPLOAD_DIR) / file_record.file_path
            
            # 提取文本
            text = await self.extract_text(file_path)
            
            # 构建分析提示
            if not query:
                query = "请对这篇文档进行总结，包括主要内容、关键点和结论。"
                
            if not system_prompt:
                system_prompt = """你是一个专业的文档分析助手。请仔细分析文档内容，提供准确、清晰的分析结果。
                分析应该包括：
                1. 文档主题和类型
                2. 主要内容概述
                3. 关键点分析
                4. 结论或建议
                请用清晰的结构化格式呈现分析结果。"""
            
            # 调用AI进行分析
            analysis_result = await ai_client.analyze_document(
                text=text,
                query=query,
                system_prompt=system_prompt
            )
            
            # 保存分析记录
            analysis_record = AnalysisRecord(
                file_id=file_id,
                analysis_type="document",
                result=analysis_result
            )
            db.add(analysis_record)
            await db.commit()
            
            return {
                "file_id": file_id,
                "original_name": file_record.original_name,
                "analysis": analysis_result
            }
            
        except Exception as e:
            app_logger.error(f"文档分析失败: {str(e)}")
            raise

    async def analyze_multiple_documents(
        self,
        db: AsyncSession,
        file_ids: List[str],
        query: Optional[str] = None,
        system_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """分析多个文档"""
        try:
            results = []
            for file_id in file_ids:
                result = await self.analyze_document(db, file_id, query, system_prompt)
                results.append(result)
            
            # 构建比较分析提示
            compare_prompt = """请对这些文档进行对比分析，包括：
            1. 共同点
            2. 差异点
            3. 综合结论
            请用清晰的结构化格式呈现分析结果。"""
            
            # 调用AI进行比较分析
            comparison_result = await ai_client.analyze_multiple_documents(
                documents=[r["analysis"] for r in results],
                query=query or "请对这些文档进行对比分析",
                system_prompt=compare_prompt
            )
            
            return {
                "individual_analyses": results,
                "comparison_analysis": comparison_result
            }
            
        except Exception as e:
            app_logger.error(f"多文档分析失败: {str(e)}")
            raise

    async def _extract_pdf_text_from_bytes(self, file_content: bytes) -> str:
        """从二进制内容中提取PDF文本"""
        try:
            # 使用线程池执行同步操作
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._pdf_extract_worker_from_bytes, file_content)
        except Exception as e:
            raise ValueError(f"PDF文本提取失败: {str(e)}")

    def _pdf_extract_worker_from_bytes(self, file_content: bytes) -> str:
        """PDF文本提取工作函数(二进制版本)"""
        text = []
        pdf = PdfReader(io.BytesIO(file_content))
        for page in pdf.pages:
            text.append(page.extract_text())
        return '\n'.join(text)

    async def extract_text_from_url(self, url: str) -> str:
        """从URL中提取文本内容"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status != 200:
                        raise ValueError(f"无法访问文档URL: {url}")
                    
                    # 读取文件内容
                    file_content = await response.read()
                    
                    # 根据URL后缀处理不同类型的文件
                    if url.lower().endswith('.docx'):
                        doc = Document(io.BytesIO(file_content))
                        paragraphs = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():  # 只添加非空段落
                                paragraphs.append(paragraph.text)
                        return '\n'.join(paragraphs)
                    else:
                        raise ValueError(f"不支持的文件类型: {url}")
                    
        except Exception as e:
            app_logger.error(f"从URL提取文本失败: {str(e)}")
            raise ValueError(f"从URL提取文本失败: {str(e)}")

# 创建文档服务实例
document_service = DocumentService()