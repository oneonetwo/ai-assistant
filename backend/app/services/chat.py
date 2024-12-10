from typing import List, Dict, Any, Optional, AsyncGenerator
from sqlalchemy.ext.asyncio import AsyncSession
from app.services.context import get_conversation, add_message, get_context_messages
from app.models.schemas import MessageCreate
from app.services.ai_client import ai_client
from app.core.config import settings
from app.core.logging import app_logger
from app.services.exceptions import NotFoundError, APIError
from sqlalchemy import select, and_, desc
from app.db.models import Message, File
import json
from pathlib import Path
import aiofiles
import aiohttp
import uuid
import io
from pypdf import PdfReader
from docx import Document
import tempfile
import subprocess
import markdown
from bs4 import BeautifulSoup
import re
import chardet

async def process_chat(
    db: AsyncSession,
    session_id: str,
    user_message: str
) -> Dict[str, str]:
    """处理普通聊天请求"""
    # 获取会话
    conversation = await get_conversation(db, session_id)
    if not conversation:
        raise NotFoundError(detail=f"会话 {session_id} 不存在")

    try:
        # 保存用户消息
        user_msg = MessageCreate(
            role="user",
            content=user_message,
            file_id=None
        )
        saved_user_msg = await add_message(db, conversation.id, user_msg)

        # 获取上下文消息
        context_messages = await get_context_messages(
            db,
            conversation.id,
            settings.MAX_CONTEXT_TURNS
        )

        # 转换为AI客户端所需格式
        messages = [
            {"role": msg.role, "content": msg.content}
            for msg in context_messages
        ]
        messages.append({"role": "user", "content": user_message})

        # 生成AI回复
        ai_response = await ai_client.generate_response(messages)

        # 保存AI回复，关联到用户消息
        ai_msg = MessageCreate(
            role="assistant",
            content=ai_response,
            parent_message_id=saved_user_msg.id  # 设置父消息ID，建立问答关系
        )
        await add_message(db, conversation.id, ai_msg)

        return {
            "session_id": session_id,
            "response": ai_response
        }
    except Exception as e:
        app_logger.error(f"处理聊天请求失败: {str(e)}")
        raise

async def initialize_stream_chat(
    db: AsyncSession,
    session_id: str,
    user_message: str
) -> None:
    """初始化流式聊天"""
    conversation = await get_conversation(db, session_id)
    if not conversation:
        raise NotFoundError(detail=f"话 {session_id} 不存在")

    try:
        # 保存用户消息
        user_msg = MessageCreate(role="user", content=user_message)
        saved_user_msg = await add_message(db, conversation.id, user_msg)

        # 获取上下文消息
        context_messages = await get_context_messages(
            db,
            conversation.id,
            settings.MAX_CONTEXT_TURNS
        )

        # 转换为AI客户端所需格式
        messages = [
            {"role": msg.role, "content": msg.content}
            for msg in context_messages
        ]
        
        # 添加当前用户消息
        messages.append({"role": "user", "content": user_message})

        # 初始化流响应
        await ai_client.initialize_stream(session_id, messages)

    except Exception as e:
        app_logger.error(f"初始化流式聊天失败: {str(e)}")
        raise APIError(detail=f"初始化流式聊天失败: {str(e)}")

async def process_stream_chat(
    db: AsyncSession,
    session_id: str,
) -> AsyncGenerator[str, None]:
    """处理流式聊天响应"""
    conversation = None
    full_response = ""
    
    try:
        conversation = await get_conversation(db, session_id)
        if not conversation:
            raise NotFoundError(detail=f"会话 {session_id} 不存在")

        # 发送开始事件
        yield f"data: {json.dumps({'type': 'start', 'data': {}}, ensure_ascii=False)}\n\n"

        # 使用新的据库会话来获取最后的用户消息
        last_user_message = await get_last_user_message(db, conversation.id)
        if not last_user_message:
            app_logger.warning(f"未找到用户消息: conversation_id={conversation.id}")
            raise APIError(detail="未找到相关的用户消息")

        async for response_chunk in ai_client.get_stream_response(session_id):
            full_response += response_chunk
            # 发送数据块事件
            chunk_data = {
                'type': 'chunk',
                'data': {
                    'content': response_chunk,
                    'full_text': full_response
                }
            }
            yield f"data: {json.dumps(chunk_data, ensure_ascii=False)}\n\n"

        # 发送结束事件
        end_data = {
            'type': 'end',
            'data': {
                'full_text': full_response
            }
        }
        yield f"data: {json.dumps(end_data, ensure_ascii=False)}\n\n"

    except Exception as e:
        app_logger.error(f"处理流式聊天失败: {str(e)}")
        error_data = {
            'type': 'error',
            'data': {
                'message': str(e)
            }
        }
        yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"
        return  # 确保在错误情况下提前返回
        
    finally:
        await ai_client.cleanup_stream(session_id)

    # 在生成响应完成后，使用新的事务保存消息
    try:
        if full_response and conversation:
            # 创建并保存AI回复消息
            ai_msg = MessageCreate(
                role="assistant",
                content=full_response,
                parent_message_id=last_user_message.id
            )
            
            # 保存消息并提交事务
            await add_message(db, conversation.id, ai_msg)
            await db.commit()
            
            app_logger.info(
                f"成功保存AI回复消息: conversation_id={conversation.id}, "
                f"parent_message_id={last_user_message.id}, "
                f"content_preview={full_response[:100]}..."
            )
    except Exception as e:
        app_logger.error(f"保存AI回复消息失败: {str(e)}")
        await db.rollback()
        # 这里我们不抛出异常，因为消息已经发送给了用��

async def get_last_user_message(
    db: AsyncSession,
    conversation_id: int
) -> Optional[Message]:
    """获取最后一条用户消息"""
    query = (
        select(Message)
        .where(
            and_(
                Message.conversation_id == conversation_id,
                Message.role == "user"
            )
        )
        .order_by(desc(Message.created_at))
        .limit(1)
    )
    result = await db.execute(query)
    return result.scalar_one_or_none()

async def init_image_stream_chat(
    db: AsyncSession,
    session_id: str,
    message: str,
    image_url: str,
    file_id: str
) -> None:
    """初始化流式图片聊天"""
    conversation = await get_conversation(db, session_id)
    if not conversation:
        raise NotFoundError(detail=f"会话 {session_id} 不存在")

    try:
        # 保存用户消息到数据库
        user_msg = MessageCreate(
            role="user",
            content=message,     # 只存储纯文本消息
            file_id=file_id      # 文件ID存储在专门的字段中
        )
        saved_user_msg = await add_message(db, conversation.id, user_msg)

        # 获取上下文消息
        context_messages = await get_context_messages(
            db,
            conversation.id,
            settings.MAX_CONTEXT_TURNS
        )

        # 转换为AI客户端所需格式
        messages = []
        for msg in context_messages[:-1]:  # 排除最后条消息
            message_data = {"role": msg.role, "content": msg.content}
            if msg.file_id:  # 如果消息关联文件，添加文件信息
                # 获取文件信息
                file_query = select(File).where(File.file_id == msg.file_id)
                result = await db.execute(file_query)
                file_record = result.scalar_one_or_none()
                if file_record and file_record.file_type == "image":
                    message_data["content"] = [
                        {"type": "text", "text": msg.content},
                        {"type": "image_url", "image_url": {"url": file_record.file_path}}
                    ]
            messages.append(message_data)

        # 添加当前用户消息，使用正确的Vision API格式
        messages.append({
            "role": "user",
            "content": [
                {"type": "text", "text": message},
                {"type": "image_url", "image_url": {"url": image_url}}
            ]
        })

        # 初始化流式响应
        await ai_client.initialize_stream(
            session_id=session_id,
            messages=messages,
            model=ai_client.vision_model
        )

    except Exception as e:
        app_logger.error(f"始化流式图片聊天失败: {str(e)}")
        raise APIError(detail=f"初始化流式图片聊天失败: {str(e)}")

async def init_file_stream_chat(
    db: AsyncSession,
    session_id: str,
    message: str,
    file_id: str,
    file_type: str,
    file_text: str
) -> None:
    """初始化流式文件聊天"""
    conversation = await get_conversation(db, session_id)
    if not conversation:
        raise NotFoundError(detail=f"会话 {session_id} 不存在")

    try:
        # 保存用户消息到数据库
        user_msg = MessageCreate(
            role="user",
            content=message,
            file_id=file_id
        )
        saved_user_msg = await add_message(db, conversation.id, user_msg)

        # 获取上下文消息
        context_messages = await get_context_messages(
            db,
            conversation.id,
            settings.MAX_CONTEXT_TURNS
        )

        # 转换为AI客户端所需格式
        messages = []
        for msg in context_messages[:-1]:  # 排除最后一条消息
            message_data = {"role": msg.role, "content": msg.content}
            messages.append(message_data)

        # 添加当前用户消息，包含文件内容
        if file_type == "image":
            file_query = select(File).where(File.file_id == file_id)
            result = await db.execute(file_query)
            file_record = result.scalar_one_or_none()
            messages.append({
                "role": "user",
                "content": [
                    {"type": "text", "text": message},
                    {"type": "image_url", "image_url": {"url": file_record.file_path}}
                ]
            })
        else:
            # 对于文档类型，将文件内容添加到消息中
            messages.append({
                "role": "user",
                "content": f"{message}\n\n文档内容：\n{file_text}"
            })

        # 初始化流式响应
        model = ai_client.vision_model if file_type == "image" else ai_client.model
        await ai_client.initialize_stream(
            session_id=session_id,
            messages=messages,
            model=model
        )

    except Exception as e:
        app_logger.error(f"初始化流式文件聊天失败: {str(e)}")
        raise APIError(detail=f"初始化流式文件聊天失败: {str(e)}")

@staticmethod
async def extract_text(file_path: str) -> str:
    """从文件中提取文本"""
    try:
        app_logger.info(f"开始提取文件文: {file_path}")
        
        # 如果是URL，尝试下载文件内容
        if file_path.startswith(('http://', 'https://')):
            app_logger.info("检测到URL文件，开始下载")
            async with aiohttp.ClientSession() as session:
                async with session.get(file_path) as response:
                    if response.status == 200:
                        file_content = await response.read()
                        app_logger.info(f"文件下载成功，大小: {len(file_content)} bytes")
                        
                        # 根据文件扩展名处理不同类型
                        file_ext = Path(file_path).suffix.lower()
                        if file_ext == '.pdf':
                            return await process_pdf_content(file_content)
                        elif file_ext in ['.docx', '.doc']:
                            return await process_word_content(file_content, file_ext)
                        elif file_ext == '.md':
                            return await process_markdown_content(file_content.decode('utf-8'))
                        elif file_ext == '.txt':
                            return await process_txt_content(file_content)
                    else:
                        error_msg = f"下载文件失败，状态码: {response.status}"
                        app_logger.error(error_msg)
                        return error_msg
        
        # 处理本地文件
        suffix = Path(file_path).suffix.lower()
        app_logger.info(f"处理本地文件，文件类型: {suffix}")
        
        if suffix == '.pdf':
            with open(file_path, 'rb') as f:
                return await process_pdf_content(f.read())
        elif suffix in ['.docx', '.doc']:
            with open(file_path, 'rb') as f:
                return await process_word_content(f.read(), suffix)
        elif suffix == '.md':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                return await process_markdown_content(content)
        elif suffix == '.txt':
            with open(file_path, 'rb') as f:
                return await process_txt_content(f.read())
        else:
            error_msg = f"不支持的文件类型: {suffix}"
            app_logger.error(error_msg)
            return error_msg

    except Exception as e:
        app_logger.error(f"文本提取失败: {str(e)}", exc_info=True)
        return f"文本提取失败: {str(e)}"

@staticmethod
async def process_pdf_content(content: bytes) -> str:
    """处理PDF文件内容"""
    try:
        app_logger.info("开始处理PDF文")
        pdf = PdfReader(io.BytesIO(content))
        text = []
        total_pages = len(pdf.pages)
        app_logger.info(f"PDF文件共 {total_pages} 页")
        
        for i, page in enumerate(pdf.pages, 1):
            app_logger.debug(f"正在处理第 {i}/{total_pages} 页")
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
                
        extracted_text = '\n'.join(text)
        app_logger.info(f"PDF文本提取成功，提取长度: {len(extracted_text)}")
        return extracted_text
    except Exception as e:
        app_logger.error(f"PDF文件处理失败: {str(e)}", exc_info=True)
        return f"PDF文件处理失败: {str(e)}"

@staticmethod
async def process_word_content(content: bytes, file_ext: str) -> str:
    """处理Word文档内容"""
    try:
        app_logger.info(f"开始处理Word文档 ({file_ext})")
        
        if file_ext == '.docx':
            # 处理.docx文件
            doc = Document(io.BytesIO(content))
            paragraphs = []
            
            # 提取段落文本
            app_logger.info(f"文档共有 {len(doc.paragraphs)} 个段落")
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            extracted_text = '\n'.join(paragraphs)
            app_logger.info(f"DOCX文本提取成功，提取长度: {len(extracted_text)}")
            return extracted_text
            
        elif file_ext == '.doc':
            # 处理.doc文件 (需要使用临时文件)
            app_logger.info("处理.doc文件，使用临时文件")
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                # 使用 antiword 转换 .doc 文件 (需要系统安装 antiword)
                result = subprocess.run(
                    ['antiword', temp_path],
                    capture_output=True,
                    text=True,
                    encoding='utf-8'
                )
                
                if result.returncode == 0:
                    extracted_text = result.stdout
                    app_logger.info(f"DOC文本提取成功，提取长度: {len(extracted_text)}")
                    return extracted_text
                else:
                    error_msg = f"DOC文件处理失败: {result.stderr}"
                    app_logger.error(error_msg)
                    return error_msg
                    
            finally:
                # 清理临时文件
                Path(temp_path).unlink(missing_ok=True)
                
    except Exception as e:
        app_logger.error(f"Word文档处理失败: {str(e)}", exc_info=True)
        return f"Word文档处理失败: {str(e)}"

@staticmethod
async def process_markdown_content(content: str) -> str:
    """处理Markdown文件内容"""
    try:
        app_logger.info("开始处理Markdown文件")
        
        # 保存原始的代码块
        code_blocks = {}
        
        def save_code_block(match):
            """保存代码块并返回占位符"""
            block_id = f"CODE_BLOCK_{len(code_blocks)}"
            code_blocks[block_id] = match.group(1)
            return block_id
        
        # 保存代码块
        content_with_placeholders = re.sub(
            r'```[\w]*\n(.*?)```',
            save_code_block,
            content,
            flags=re.DOTALL
        )
        
        # 转换Markdown为HTML
        html = markdown.markdown(
            content_with_placeholders,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.nl2br'
            ]
        )
        
        # 使用BeautifulSoup提取文本
        soup = BeautifulSoup(html, 'html.parser')
        
        # 提取所有文本内容
        text_parts = []
        
        # 处理标题
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(heading.name[1])
            text_parts.append(f"{'#' * level} {heading.get_text().strip()}\n")
        
        # 处理段落
        for p in soup.find_all('p'):
            text_parts.append(p.get_text().strip() + '\n')
        
        # 处理列表
        for ul in soup.find_all(['ul', 'ol']):
            for li in ul.find_all('li'):
                text_parts.append(f"- {li.get_text().strip()}\n")
        
        # 处理表格
        for table in soup.find_all('table'):
            for row in table.find_all('tr'):
                cells = [cell.get_text().strip() for cell in row.find_all(['th', 'td'])]
                text_parts.append(' | '.join(cells) + '\n')
        
        # 合并文本
        extracted_text = '\n'.join(text_parts)
        
        # 还原代码块
        for block_id, code in code_blocks.items():
            extracted_text = extracted_text.replace(block_id, f"\n```\n{code}\n```\n")
        
        # 清理多余的空行
        cleaned_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
        
        app_logger.info(f"Markdown文本提取成功，提取长度: {len(cleaned_text)}")
        
        # 添加文档结构信息
        structure_info = analyze_markdown_structure(content)
        app_logger.info("Markdown文档结构分析完成")
        
        # 返回处理后的文本，包含文档结构信息
        final_text = (
            "# 文档结构信息\n"
            f"{structure_info}\n\n"
            "# 文档内容\n"
            f"{cleaned_text}"
        )
        
        return final_text
        
    except Exception as e:
        app_logger.error(f"Markdown文档处理失败: {str(e)}", exc_info=True)
        return f"Markdown文档处理失败: {str(e)}"

@staticmethod
def analyze_markdown_structure(content: str) -> str:
    """分析Markdown文档结构"""
    try:
        # 提取所有标题
        headers = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
        
        # 统计代码块
        code_blocks = re.findall(r'```(\w*)', content)
        code_languages = [lang for lang in code_blocks if lang]
        
        # 统计链接和图片
        links = re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', content)
        images = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', content)
        
        # 构建结构信息
        structure = []
        structure.append("## 文档大纲")
        for level, title in headers:
            indent = "  " * (len(level) - 1)
            structure.append(f"{indent}- {title}")
        
        if code_languages:
            structure.append("\n## 代码块信息")
            lang_count = {}
            for lang in code_languages:
                lang_count[lang] = lang_count.get(lang, 0) + 1
            for lang, count in lang_count.items():
                lang_name = lang if lang else '未指定语言'
                structure.append(f"- {lang_name}: {count} 个代码块")
        
        if links:
            structure.append("\n## 链接信息")
            structure.append(f"- 总计: {len(links)} 个链接")
        
        if images:
            structure.append("\n## 图片信息")
            structure.append(f"- 总计: {len(images)} 张图片")
        
        return '\n'.join(structure)
        
    except Exception as e:
        app_logger.error(f"Markdown结构分析失败: {str(e)}", exc_info=True)
        return "Markdown结构分析失败"

async def initialize_message_analysis_stream(
    messages: List[Dict[str, Any]],
    system_prompt: Optional[str] = None
) -> str:
    """初始化消息分析流
    
    Args:
        messages: 消息列表，每个消息包含 role 和 content
        system_prompt: 可选的系统提示
    
    Returns:
        str: 分析会话ID
    """
    try:
        print("process_message_analysis_stream***********************")
        # 构建分析提示
        message_texts = []
        for msg in messages:
            # 使用字典访问语法获取role和content
            formatted_msg = f"{msg['role']}: {msg['content']}"
            message_texts.append(formatted_msg)
            
        combined_message = "\n".join(message_texts)
        analysis_prompt = f"""请分析以下对话内容:
{combined_message}

请提供:
1. 一个不超过30个字的精简标题
2. 提取有效的关键字
3. 提取考点，要点, 简短
4. 请分析内容，结构清晰，内容完整

输出格式:
标题: <标题>

关键字: <关键字>

要点: <要点>

 <分析内容>"""

        # 构建消息列表
        prompt_messages = []
        if system_prompt:
            prompt_messages.append({"role": "system", "content": system_prompt})
        prompt_messages.append({"role": "user", "content": analysis_prompt})
        # 初始化分析流并返回会话ID
        session_id = await ai_client.initialize_analysis_stream(
            messages=prompt_messages,
            temperature=0.7
        )
        app_logger.info(f"成功初始化消息分析流，会话ID: {session_id}")
        return session_id
    except Exception as e:
        app_logger.error(f"初始化消息分析流失败: {str(e)}")
        raise APIError(f"初始化消息分析流失败: {str(e)}")

async def get_message_analysis_stream(session_id: str) -> AsyncGenerator[str, None]:
    """获取消息分析的流式响应"""
    try:
        # 发送开始事件
        yield f"data: {json.dumps({'type': 'start', 'data': {}}, ensure_ascii=False)}\n\n"
        
        # 检查会话是否存在
        is_active = await ai_client.is_session_active(session_id)
        if not is_active:
            error_data = {
                'type': 'error',
                'data': {
                    'message': '会话未初始化或已过期'
                }
            }
            yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"
            return

        full_response = ""
        async for chunk in ai_client.get_analysis_stream(session_id):
            # 确保chunk是UTF-8编码
            if isinstance(chunk, bytes):
                chunk = chunk.decode('utf-8')
            
            full_response += chunk
            # 构建数据块事件并确保使用UTF-8编码
            chunk_data = {
                'type': 'chunk',
                'data': {
                    'content': chunk,
                    'section': 'analysis',
                    'full_text': full_response
                }
            }
            yield f"data: {json.dumps(chunk_data, ensure_ascii=False)}\n\n".encode('utf-8').decode('utf-8')

        # 发送结束事件
        end_data = {
            'type': 'end',
            'data': {
                'full_text': full_response
            }
        }
        yield f"data: {json.dumps(end_data, ensure_ascii=False)}\n\n".encode('utf-8').decode('utf-8')

    except Exception as e:
        app_logger.error(f"获取消息分析流失败: {str(e)}")
        error_data = {
            'type': 'error',
            'data': {
                'message': str(e)
            }
        }
        yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"

@staticmethod
async def process_txt_content(content: bytes) -> str:
    """处理TXT文件内容"""
    try:
        app_logger.info("开始处理TXT文件")
        app_logger.info(f"文件大小: {len(content)} bytes")

        # 检测文件编码
        detection = chardet.detect(content)
        encoding = detection['encoding']
        confidence = detection['confidence']
        app_logger.info(f"检测到编码: {encoding} (置信度: {confidence})")

        # 尝试不同的编码方式
        encodings_to_try = [
            encoding,  # 检测到的编码
            'utf-8',
            'gb18030',  # 支持中文
            'gbk',
            'gb2312',
            'ascii',
            'iso-8859-1'
        ]

        text = None
        used_encoding = None

        for enc in encodings_to_try:
            if not enc:
                continue
            try:
                text = content.decode(enc)
                used_encoding = enc
                app_logger.info(f"成功使用 {enc} 编码解析文件")
                break
            except UnicodeDecodeError:
                app_logger.debug(f"使用 {enc} 编码解析失败，尝试下一个编码")
                continue

        if text is None:
            error_msg = "无法使用任何编码解析文件内容"
            app_logger.error(error_msg)
            return error_msg

        # 处理文本内容
        # 1. 规范化���行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 2. 删除连续的空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 3. 删除行尾空白字符
        text = '\n'.join(line.rstrip() for line in text.split('\n'))

        # 4. 确保文件开头和结尾没有多余的空行
        text = text.strip()

        # 计算行数
        line_count = text.count('\n') + 1

        # 构建文件信息字符串
        info_parts = [
            "# 文件信息",
            f"- 文件大小: {len(content)} bytes",
            f"- 使用编码: {used_encoding}",
            f"- 行数: {line_count}",
            f"- 字符数: {len(text)}",
            "",
            "# 文件内容",
            ""
        ]
        
        file_info = '\n'.join(info_parts)
        final_text = f"{file_info}{text}"

        # 构建日志信息字符串
        log_parts = [
            "TXT文件处理完成:",
            f"- 最终使用的编码: {used_encoding}",
            f"- 处理后文本长度: {len(final_text)}",
            f"- 行数: {line_count}"
        ]
        
        app_logger.info('\n'.join(log_parts))

        return final_text

    except Exception as e:
        app_logger.error(f"TXT文件处理失败: {str(e)}", exc_info=True)
        return f"TXT文件处理失败: {str(e)}"